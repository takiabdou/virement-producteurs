"""
Script de génération du guide d'utilisateur
Application Virement Producteurs - CRMA
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from datetime import datetime
import os

def creer_guide_utilisateur():
    """Crée le guide d'utilisateur au format DOCX"""
    
    doc = Document()
    
    # Style pour les titres
    styles = doc.styles
    
    # Titre principal
    titre = doc.add_heading('GUIDE D\'UTILISATION', 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sous-titre
    sous_titre = doc.add_heading('Application Virement Producteurs', level=1)
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informations CRMA
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Caisses Régionales de Mutualité Agricole (CRMA)\n')
    p.add_run('République Algérienne Démocratique et Populaire\n')
    p.add_run(f'\nVersion 1.0 - Juillet 2026')
    
    doc.add_paragraph('\n' * 2)
    
    # Introduction
    doc.add_heading('À PROPOS DE CE GUIDE', level=1)
    doc.add_paragraph(
        'Ce guide est conçu pour les utilisateurs n\'ayant aucune connaissance en informatique. '
        'Chaque étape est expliquée de manière simple et détaillée avec des exemples concrets.'
    )
    
    # Table des matières
    doc.add_page_break()
    doc.add_heading('SOMMAIRE', level=1)
    
    sommaire = [
        '1. Introduction',
        '2. Accès à l\'application',
        '3. La page de connexion',
        '4. Le tableau de bord',
        '5. Le brouillard de caisse',
        '6. Éditer un bon de versement',
        '7. Consulter l\'historique',
        '8. Questions fréquentes',
        '9. Assistance et contact'
    ]
    
    for item in sommaire:
        doc.add_paragraph(item)
    
    # Chapitre 1
    doc.add_page_break()
    doc.add_heading('CHAPITRE 1 : INTRODUCTION', level=1)
    
    doc.add_heading('1.1 Qu\'est-ce que l\'application Virement Producteurs ?', level=2)
    doc.add_paragraph(
        'L\'application Virement Producteurs est un outil informatique qui permet de :'
    )
    doc.add_paragraph(
        '• Enregistrer les encaissements journaliers de votre Bureau Local (BL)\n'
        '• Calculer automatiquement les montants à verser\n'
        '• Générer les bons de versement (CCP, BADR, BNA)\n'
        '• Consulter l\'historique des versements',
        style='List Bullet'
    )
    
    doc.add_heading('1.2 À qui s\'adresse cette application ?', level=2)
    doc.add_paragraph(
        'Cette application est destinée aux agents des Bureaux Locaux des CRMA qui sont '
        'chargés de la saisie des encaissements, de la préparation des versements au siège '
        'et de la gestion des bons de paiement.'
    )
    
    doc.add_heading('1.3 Ce dont vous avez besoin', level=2)
    doc.add_paragraph(
        'Matériel nécessaire :\n'
        '• Un ordinateur (Windows 7, 10 ou 11)\n'
        '• Une connexion au réseau de la CRMA\n'
        '• Vos identifiants de connexion (nom d\'utilisateur et mot de passe)',
        style='List Bullet'
    )
    
    # Chapitre 2
    doc.add_page_break()
    doc.add_heading('CHAPITRE 2 : ACCÈS À L\'APPLICATION', level=1)
    
    doc.add_heading('2.1 Démarrer l\'application', level=2)
    
    doc.add_paragraph('Étape 1 : Double-cliquez sur l\'icône "demarrer.bat" sur votre bureau')
    doc.add_paragraph('Étape 2 : Une fenêtre bleue s\'ouvre automatiquement')
    doc.add_paragraph('Étape 3 : Attendez quelques secondes - votre navigateur Internet s\'ouvre automatiquement')
    doc.add_paragraph('Étape 4 : La page de connexion de l\'application apparaît')
    
    # Encadré information
    p = doc.add_paragraph()
    p.add_run('Information importante : ').bold = True
    p.add_run('Ne fermez pas la fenêtre bleue pendant que vous utilisez l\'application. '
              'Si vous la fermez, l\'application ne fonctionnera plus.')
    
    # Chapitre 3
    doc.add_page_break()
    doc.add_heading('CHAPITRE 3 : LA PAGE DE CONNEXION', level=1)
    
    doc.add_heading('3.1 Se connecter', level=2)
    doc.add_paragraph(
        'Sur la page de connexion, vous voyez :\n'
        '• Le logo CNMA en haut\n'
        '• Un cadre blanc avec deux cases à remplir\n'
        '• Un bouton bleu "Se connecter"',
        style='List Bullet'
    )
    
    doc.add_paragraph('Étape 1 : Cliquez dans la première case "Nom d\'utilisateur"')
    doc.add_paragraph('Étape 2 : Tapez votre nom d\'utilisateur (exemple : user001)')
    doc.add_paragraph('Étape 3 : Cliquez dans la deuxième case "Mot de passe"')
    doc.add_paragraph('Étape 4 : Tapez votre mot de passe (les caractères ne s\'affichent pas, c\'est normal)')
    doc.add_paragraph('Étape 5 : Cliquez sur le bouton bleu "Se connecter →"')
    
    # Avertissement
    p = doc.add_paragraph()
    p.add_run('Attention : ').bold = True
    p.add_run('Si vous voyez un message d\'erreur rouge, vérifiez :\n'
              '• Que votre nom d\'utilisateur est correct\n'
              '• Que votre mot de passe est correct\n'
              '• Que la touche Verr Maj (Majuscule) n\'est pas activée')
    
    doc.add_heading('3.2 Que faire si j\'ai oublié mon mot de passe ?', level=2)
    doc.add_paragraph(
        'Contactez votre responsable ou le service informatique de votre CRMA pour '
        'réinitialiser votre mot de passe.'
    )
    
    # Chapitre 4
    doc.add_page_break()
    doc.add_heading('CHAPITRE 4 : LE TABLEAU DE BORD', level=1)
    
    doc.add_heading('4.1 Découverte du tableau de bord', level=2)
    doc.add_paragraph(
        'Après la connexion, vous arrivez sur le tableau de bord. C\'est la page principale '
        'de l\'application.'
    )
    doc.add_paragraph(
        'Vous voyez :\n'
        '• En haut : une barre bleue avec le nom de l\'application et votre nom\n'
        '• Un message de bienvenue\n'
        '• Trois grandes cases (boutons) pour accéder aux fonctions principales',
        style='List Bullet'
    )
    
    doc.add_heading('4.2 Les trois fonctions principales', level=2)
    
    doc.add_paragraph('1. Brouillard de caisse (case avec icône )')
    doc.add_paragraph('   C\'est ici que vous enregistrez les encaissements de la journée')
    
    doc.add_paragraph('2. Éditer un versement (case avec icône 💳)')
    doc.add_paragraph('   C\'est ici que vous créez les bons de versement (CCP, BADR, BNA)')
    
    doc.add_paragraph('3. Historique (case avec icône )')
    doc.add_paragraph('   C\'est ici que vous consultez les bons déjà créés')
    
    doc.add_heading('4.3 La barre de navigation', level=2)
    doc.add_paragraph(
        'En haut de chaque page, vous trouvez une barre bleue avec :\n'
        '• 🏠 Accueil : Retour au tableau de bord\n'
        '•  Brouillard : Accéder au brouillard de caisse\n'
        '• 📋 Historique : Consulter l\'historique\n'
        '• ⏏ Déconnexion (en rouge) : Pour quitter l\'application',
        style='List Bullet'
    )
    
    # Chapitre 5
    doc.add_page_break()
    doc.add_heading('CHAPITRE 5 : LE BROUILLARD DE CAISSE', level=1)
    
    doc.add_heading('5.1 Qu\'est-ce que le brouillard de caisse ?', level=2)
    doc.add_paragraph(
        'Le brouillard de caisse est un registre où vous notez tous les encaissements de '
        'la journée. C\'est comme un cahier de recettes quotidien.'
    )
    
    doc.add_heading('5.2 Accéder au brouillard', level=2)
    doc.add_paragraph(
        'Depuis le tableau de bord, cliquez sur la case "Brouillard de caisse"\n'
        'OU\n'
        'Cliquez sur "Brouillard" dans la barre de navigation en haut'
    )
    
    doc.add_heading('5.3 Ajouter un encaissement', level=2)
    doc.add_paragraph(
        'En bas de la page du brouillard, vous voyez un formulaire avec deux cases :'
    )
    
    doc.add_paragraph('Étape 1 : Dans la case "N° Contrat (optionnel)", tapez le numéro de contrat si vous en avez un')
    doc.add_paragraph('   Exemple : 2026/12345')
    doc.add_paragraph('Étape 2 : Dans la case "Montant encaissé (DA)", tapez le montant reçu')
    doc.add_paragraph('   Exemple : 15000 ou 15000,50')
    doc.add_paragraph('Étape 3 : Cliquez sur le bouton vert "+ Ajouter"')
    doc.add_paragraph('Résultat : La ligne apparaît dans le tableau avec l\'heure et le numéro')
    
    doc.add_paragraph(
        'Important :\n'
        '• Vous pouvez utiliser le point (.) ou la virgule (,) pour les décimales\n'
        '• Le numéro de contrat est facultatif - vous pouvez le laisser vide\n'
        '• Chaque ligne est automatiquement datée et horodatée',
        style='List Bullet'
    )
    
    doc.add_heading('5.4 Consulter le brouillard', level=2)
    
    # Tableau explicatif
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Colonne'
    hdr_cells[1].text = 'Ce qu\'elle contient'
    
    donnees = [
        ('N°', 'Numéro de la ligne (1, 2, 3...)'),
        ('Heure', 'L\'heure de l\'enregistrement'),
        ('N° Contrat', 'Le numéro de contrat si renseigné'),
        ('Montant (DA)', 'Le montant encaissé en dinars'),
        ('Saisi par', 'Le nom de la personne qui a saisi')
    ]
    
    for col, contenu in donnees:
        row_cells = table.add_row().cells
        row_cells[0].text = col
        row_cells[1].text = contenu
    
    doc.add_paragraph('En bas du tableau, vous voyez le TOTAL DU JOUR qui s\'affiche en bleu.')
    
    doc.add_heading('5.5 Modifier ou supprimer une ligne', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Important : ').bold = True
    p.add_run('Vous ne pouvez supprimer que les lignes que VOUS avez saisies aujourd\'hui. '
              'Vous ne pouvez pas supprimer les lignes saisies par d\'autres agents.')
    
    doc.add_paragraph('Pour supprimer une ligne que vous avez saisie :')
    doc.add_paragraph('Étape 1 : Trouvez la ligne à supprimer dans le tableau')
    doc.add_paragraph('Étape 2 : Cliquez sur le bouton rouge ✕ à droite de la ligne')
    doc.add_paragraph('Étape 3 : Une question apparaît : "Supprimer cette ligne ?"')
    doc.add_paragraph('Étape 4 : Cliquez sur "OK" pour confirmer ou "Annuler"')
    
    doc.add_heading('5.6 Consulter un jour précédent', level=2)
    doc.add_paragraph(
        'En haut du brouillard, vous voyez une section "Consulter un brouillard" :'
    )
    doc.add_paragraph('Étape 1 : Cliquez sur le calendrier ou la case date')
    doc.add_paragraph('Étape 2 : Choisissez la date que vous voulez consulter')
    doc.add_paragraph('Étape 3 : Cliquez sur le bouton bleu "Afficher"')
    doc.add_paragraph('Résultat : Le brouillard de cette date s\'affiche')
    
    doc.add_heading('5.7 Imprimer le brouillard', level=2)
    doc.add_paragraph('Étape 1 : En haut du brouillard, cliquez sur le bouton bleu "🖨️ Imprimer le brouillard"')
    doc.add_paragraph('Étape 2 : La fenêtre d\'impression de votre ordinateur s\'ouvre')
    doc.add_paragraph('Étape 3 : Choisissez votre imprimante')
    doc.add_paragraph('Étape 4 : Cliquez sur "Imprimer"')
    
    # Chapitre 6
    doc.add_page_break()
    doc.add_heading('CHAPITRE 6 : ÉDITER UN BON DE VERSEMENT', level=1)
    
    doc.add_heading('6.1 Qu\'est-ce qu\'un bon de versement ?', level=2)
    doc.add_paragraph(
        'Un bon de versement est un document officiel qui permet de transférer l\'argent '
        'encaissé vers le siège de la CRMA. Il existe trois types de versements :'
    )
    doc.add_paragraph(
        '• CCP : Versement par chèque postal (formulaire SFP 01)\n'
        '• BADR : Virement bancaire BADR\n'
        '• BNA : Virement bancaire BNA',
        style='List Bullet'
    )
    
    doc.add_heading('6.2 Accéder à la page de versement', level=2)
    doc.add_paragraph('Depuis le tableau de bord, cliquez sur "Éditer un versement"')
    
    doc.add_heading('6.3 Choisir le type de versement', level=2)
    doc.add_paragraph(
        'Sur cette page, vous voyez :\n'
        '• Le total encaissé du jour\n'
        '• Le net CCP (après déduction des droits Algérie Poste)\n'
        '• Le net BADR/BNA (sans déduction)\n'
        '• Trois boutons pour choisir le mode de versement',
        style='List Bullet'
    )
    
    # Chapitre 7
    doc.add_page_break()
    doc.add_heading('CHAPITRE 7 : CONSULTER L\'HISTORIQUE', level=1)
    
    doc.add_heading('7.1 Accéder à l\'historique', level=2)
    doc.add_paragraph(
        'Depuis le tableau de bord, cliquez sur "Historique"\n'
        'OU\n'
        'Depuis n\'importe quelle page, cliquez sur "📋 Historique" dans la barre de navigation'
    )
    
    doc.add_heading('7.2 Ce que vous voyez', level=2)
    doc.add_paragraph('Un tableau affiche tous les bons de versement déjà créés avec :')
    
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Table Grid'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Information'
    hdr_cells2[1].text = 'Description'
    
    donnees2 = [
        ('N° Émission', 'Numéro unique du bon'),
        ('Date', 'Date de création du bon'),
        ('Type', 'CCP, BADR ou BNA'),
        ('Bureau Local', 'Nom de votre BL'),
        ('Émis par', 'Nom de la personne qui a créé le bon'),
        ('Montant versé', 'Le montant en dinars')
    ]
    
    for info, desc in donnees2:
        row_cells = table2.add_row().cells
        row_cells[0].text = info
        row_cells[1].text = desc
    
    doc.add_heading('7.3 Filtrer l\'historique', level=2)
    doc.add_paragraph('En haut du tableau, vous pouvez filtrer les résultats :')
    
    doc.add_paragraph('Pour filtrer par date :')
    doc.add_paragraph('Étape 1 : Cliquez sur la case "Du" et choisissez la date de début')
    doc.add_paragraph('Étape 2 : Cliquez sur la case "Au" et choisissez la date de fin')
    doc.add_paragraph('Étape 3 : Cliquez sur "Filtrer"')
    
    doc.add_paragraph('Pour filtrer par type :')
    doc.add_paragraph('Étape 1 : Cliquez sur la liste déroulante "Type"')
    doc.add_paragraph('Étape 2 : Choisissez CCP, BADR ou BNA')
    doc.add_paragraph('Étape 3 : Cliquez sur "Filtrer"')
    
    doc.add_heading('7.4 Réimprimer un bon', level=2)
    doc.add_paragraph('Si vous avez perdu un bon ou si vous avez besoin d\'une copie :')
    doc.add_paragraph('Étape 1 : Trouvez le bon dans la liste de l\'historique')
    doc.add_paragraph('Étape 2 : Cliquez sur le bouton bleu "🖨️ Réimprimer"')
    doc.add_paragraph('Étape 3 : Le bon s\'affiche à l\'écran')
    doc.add_paragraph('Étape 4 : Cliquez sur "🖨️ Imprimer"')
    
    # Chapitre 8
    doc.add_page_break()
    doc.add_heading('CHAPITRE 8 : QUESTIONS FRÉQUENTES', level=1)
    
    doc.add_heading('8.1 Problèmes de connexion', level=2)
    doc.add_paragraph('Question : Je ne peux pas me connecter, que faire ?')
    doc.add_paragraph(
        'Réponse :\n'
        '• Vérifiez que la fenêtre bleue "demarrer.bat" est ouverte\n'
        '• Vérifiez votre nom d\'utilisateur et mot de passe\n'
        '• Vérifiez que la touche Verr Maj n\'est pas activée\n'
        '• Si le problème persiste, contactez le service informatique',
        style='List Bullet'
    )
    
    doc.add_heading('8.2 Problèmes de saisie', level=2)
    doc.add_paragraph('Question : J\'ai fait une erreur de saisie, comment corriger ?')
    doc.add_paragraph(
        'Réponse :\n'
        '• Si c\'est une ligne du brouillard que vous avez saisie aujourd\'hui : cliquez sur  pour la supprimer et ressaisissez-la\n'
        '• Si c\'est un bon déjà imprimé : contactez votre responsable\n'
        '• Vous ne pouvez pas modifier les lignes saisies par d\'autres agents',
        style='List Bullet'
    )
    
    doc.add_heading('8.3 Problèmes d\'impression', level=2)
    doc.add_paragraph('Question : L\'impression n\'est pas alignée sur le formulaire')
    doc.add_paragraph(
        'Réponse :\n'
        '• Utilisez le mode calibrage (voir section 6.5)\n'
        '• Vérifiez que vous utilisez le bon type de papier (SFP 01 pour CCP)\n'
        '• Vérifiez les marges de l\'imprimante',
        style='List Bullet'
    )
    
    # Chapitre 9
    doc.add_page_break()
    doc.add_heading('CHAPITRE 9 : ASSISTANCE ET CONTACT', level=1)
    
    doc.add_heading('9.1 En cas de problème', level=2)
    doc.add_paragraph(
        'Si vous rencontrez un problème que vous ne pouvez pas résoudre avec ce guide :'
    )
    
    p = doc.add_paragraph()
    p.add_run('Service Informatique CRMA de Saïda\n').bold = True
    p.add_run('Responsable : ').bold = True
    p.add_run('HALIMI Abdellah Takieddine\n')
    p.add_run('Fonction : Chargé du service informatique\n')
    p.add_run('Contact : Contactez votre responsable hiérarchique qui transmettra au service informatique')
    
    doc.add_heading('9.2 Avant de contacter le support', level=2)
    doc.add_paragraph(
        'Pour nous aider à résoudre votre problème rapidement, notez :\n'
        '• Le message d\'erreur exact (s\'il y en a un)\n'
        '• Ce que vous faisiez quand le problème est apparu\n'
        '• Votre Bureau Local\n'
        '• Votre nom d\'utilisateur',
        style='List Bullet'
    )
    
    # Pied de page
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('\n\nApplication Virement Producteurs\n').bold = True
    p.add_run('Développée par HALIMI Abdellah Takieddine\n')
    p.add_run('Chargé du service informatique - CRMA de Saïda\n')
    p.add_run('© 2026 Caisse Régionale de Mutualité Agricole de Saïda\n')
    p.add_run(f'\nDocument généré le {datetime.now().strftime("%d/%m/%Y à %H:%M")}')
    
    # Sauvegarder le document
    doc.save('Guide_Utilisateur_Virement_Producteurs.docx')
    print("✓ Guide DOCX créé avec succès : Guide_Utilisateur_Virement_Producteurs.docx")
    
    return True

if __name__ == "__main__":
    print("Génération du guide d'utilisateur...")
    print("=" * 60)
    creer_guide_utilisateur()
    print("=" * 60)
    print("Guide généré avec succès !")
    print("\nPour convertir en PDF, vous pouvez :")
    print("1. Ouvrir le fichier DOCX dans Microsoft Word")
    print("2. Cliquer sur Fichier > Enregistrer sous > PDF")
    print("\nOU utiliser un convertisseur en ligne gratuit")